from docx import Document
from docx.shared import Cm
from datetime import datetime

# Luo uusi Word-dokumentti
doc = Document()

# Muodosta päivämäärä
paivamaara = datetime.now().strftime("%d.%m.%Y")

# Hae ylätunniste
section = doc.sections[0]
header = section.header

# Lisää kappale ylätunnisteeseen
header_paragraph = header.paragraphs[0]

# Lisää logo
run = header_paragraph.add_run()
run.add_picture('1.png', width=Cm(3))

# Lisää päivämäärä logon viereen
header_paragraph.add_run(f'  {paivamaara}')

# Lisää otsikko
doc.add_heading('Esimerkkidokumentti', level=1)

# Lisää sisältöä
doc.add_paragraph('Tämä on Pythonilla luotu Word-tiedosto.')

# Tallenna
doc.save('esimerkki.docx')

print("Word-tiedosto luotu: esimerkki.docx")